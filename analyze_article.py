#!/usr/bin/env python
# coding: utf-8

import os, glob, re, string
# from datetime import date

import nltk
nltk.download('stopwords')
nltk.download('wordnet')
from nltk.corpus import stopwords

import pandas as pd
from pandas.plotting import table
import matplotlib.pyplot as plt
from mycolorpy import colorlist as mcp

import numpy as np
from PIL import Image
from wordcloud import WordCloud, STOPWORDS, ImageColorGenerator

import docx
from docx.shared import Inches


# Define function to load article, clean the structure and text from unwanted characters and extract seperate words
def clean_data_init(text):
    text = open(text, encoding="utf-8").read()
    lower_text = text.strip().lower()
    del_num = re.sub(r'\d+',' ',lower_text)
    
    charlist_remove = ["\t",",",".","\"","?","- ",":","’",'”','“']
    charlist_replace = ["\n","-"]
    
    remove_char = del_num
    for x in charlist_remove:
        remove_char = remove_char.replace(x,"")
        
    replace_char_space = remove_char
    for y in charlist_replace:
        replace_char_space = replace_char_space.replace(y," ")


    remove_double_space = re.sub(' +',' ',replace_char_space)
    text_csv = remove_double_space.replace(' ',',') 
    tokens = nltk.word_tokenize(remove_double_space, language='english', preserve_line=True)
    return tokens



# Define function to remove english stopwords (commonly used words which carry no substantive information)
def remove_stopword(tokens):
    custom_stop_words = {"'s","'m","'d","'t","n't"}
    stop_words = set(stopwords.words('english'))
    stop_words.update(custom_stop_words)
    filter_words = [word for word in tokens if word not in stop_words]
    return filter_words



#  Definae function to lemmatize text
def lemmatization(tokens):
    lemmatizer = nltk.WordNetLemmatizer()
    lemmatize_words = [lemmatizer.lemmatize(token) for token in tokens]
    lemmatize_words = [lemmatizer.lemmatize(token, "v") for token in lemmatize_words]
    lemmatize_words = [lemmatizer.lemmatize(token, "a") for token in lemmatize_words]
    lemmatize_words = [lemmatizer.lemmatize(token, "r") for token in lemmatize_words]
    lemmatize_words = [lemmatizer.lemmatize(token, "s") for token in lemmatize_words]
    return lemmatize_words



# Define final function to clean article's text and transform it to a list of words
def clean_data(text):
    tokens = clean_data_init(text)
    filter_words = remove_stopword(tokens)
    lemmatize_words = lemmatization(filter_words)
    final_text = ','.join(lemmatize_words)
    return final_text


# Clean article's text, get file from cwd
cwd=os.getcwd().replace('\\','\\\\')
article_text_path = glob.glob(f'{cwd}\\*.txt')[0]

data = clean_data(article_text_path)
# data

# Convert list of words to a dataframe
list = data.split(',')
df = pd.DataFrame(list, columns=['words'])
print(df)


# Check first 5 words
df_head = df["words"].head()
print(df_head)

# Count number of words in the dataframe
df_count = df.count()
print(df_count)

# Generate dataframe's basic descriptive statistics
df_descriptive_stats = df.describe()
df_descriptive_stats_string = df_descriptive_stats[['words']].to_string()
print(df_descriptive_stats)


# Get more statistics info about dataframe and words from given article
sentence_1 = "There are {} rows and {} column(s) in this dataframe.".format(df.shape[0],df.shape[1])
sentence_2 = "There are {} unique words in this article, e.g.: {}.".format(len(df.words.unique()),", ".join(df.words.unique()[0:4]))
print(sentence_1)
print(sentence_2)


# Count frequency of words 
counts_by_word = df['words'].value_counts().head(20)
print(counts_by_word)


# Present frequency of words on a chart
top_words = df.groupby('words').size().sort_values(ascending=False).head(10)

y_list = []
for y in top_words:
    y_list.append(y)

x_list = []
for x in top_words.index:
    x_list.append(x)

fig = plt.figure(figsize=(12,8))

color_map = mcp.gen_color(cmap="YlOrBr", n=10)
chart = plt.bar(x=x_list, height=y_list, color=color_map)

plt.xlabel("word")
plt.ylabel("frequency")
plt.title("Frequency of words")
chart = plt.bar_label(chart, labels=y_list, label_type="edge")
plt.show()


# Set word cloud parameters
image_mask = np.array(Image.open(glob.glob(f'{cwd}\\mask_image*')[0]))
wordcloud = WordCloud(font_path = r'C:\Windows\Fonts\Verdana.ttf',
                            stopwords = STOPWORDS,
                            background_color = 'black',
                            mask = image_mask,
                            contour_width = 2,
                            width=1584, height=396,
                            max_font_size=80, min_font_size=10,
                            colormap=plt.cm.seismic
                            # color_func=random_color_func
                            ).generate(data)


# Plot a word cloud image preview
plt.imshow(wordcloud, interpolation="bilinear")
plt.axis('off')
plt.margins(x=10, y=30)
plt.show()


# Get article's title
article_title = os.path.basename(article_text_path).split('/')[-1].split('_.')[0]
print(f"The article's title is: {article_title}")


# Save a word cloud image as png file
wordcloud.to_file(f'{article_title}_wordcloud.png')
fig.savefig(f'{article_title}_freqWords_chart.png')


# Save analysis results in a docx file
document = docx.Document()
article_title_cleaned = article_title.replace('_',' ')

document.add_heading('Text Analysis Results', 0)
document.add_heading(f'Article title: {article_title_cleaned}', 3)

h4 = document.add_heading('',4)
runner = h4.add_run('Insights:')
runner.italic = False
runner.bold = False

document.add_paragraph(sentence_1)
document.add_paragraph(sentence_2)

document.add_paragraph('Descriptive statistics:')
document.add_paragraph('#' + df_descriptive_stats_string)

document.add_picture(glob.glob(f'{cwd}\\*_freqWords_chart*')[0], width=Inches(4))
document.add_picture(glob.glob(f'{cwd}\\*_wordcloud*')[0], width=Inches(4))

document.save("article_analysis_results.docx")